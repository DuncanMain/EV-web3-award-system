from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("NEVERFLAT_TRL7_Roadmap.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(28, 35, 45)
MUTED = RGBColor(91, 101, 112)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "B7C3D0"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold=False, color=INK, size=9.5, fill: str | None = None) -> None:
    if fill:
        set_cell_shading(cell, fill)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color


def set_table_width(table, widths_inches: list[float]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    widths_dxa = [int(w * 1440) for w in widths_inches]
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_para(doc: Document, text: str = "", style: str | None = None, bold=False, color=INK, size=11, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.size = Pt(size)
        run.font.color.rgb = color
    return p


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(10.5)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.color.rgb = INK
    r2.font.size = Pt(10.5)
    r2.font.name = "Calibri"
    r2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    add_para(doc, "", after=4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        set_cell_text(hdr[idx], text, bold=True, color=DARK_BLUE, size=9.2, fill=LIGHT_FILL)
    for row_values in rows:
        row = table.add_row().cells
        for idx, text in enumerate(row_values):
            set_cell_text(row[idx], text, size=9.1)
    add_para(doc, "", after=6)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("NEVERFLAT Award and Wallet System | TRL 7 Roadmap")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("Prepared for Horizon TRL 7 readiness planning")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED


def add_title_block(doc: Document) -> None:
    add_para(doc, "", after=10)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("TRL 7 Roadmap")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("NEVERFLAT Award and Wallet System")
    r.font.size = Pt(14)
    r.font.color.rgb = MUTED
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    metadata = [
        ["Purpose", "Prioritised roadmap from current TRL 6 prototype evidence toward a credible TRL 7 claim."],
        ["Current judgement", "Estimated TRL 6. TRL 7 should be claimed only after pilot evidence, operational controls, and settlement proof are in place."],
        ["Prepared date", "6 July 2026"],
        ["Roadmap stance", "Start with stable work owned by NEVERFLAT: signed spend receipts, auditability, pilot access controls, transaction states, and evidence capture."],
    ]
    add_table(doc, ["Field", "Roadmap basis"], metadata, [1.45, 5.05])


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    add_title_block(doc)

    add_callout(
        doc,
        "Executive recommendation",
        "Do not start by rebuilding the AU ingestion path while that integration shape may still change. Start with spend receipts, audit logs, pilot access controls, and durable transaction states. Those foundations are needed regardless of how CDR ingestion finally lands.",
    )

    add_heading(doc, "1. Roadmap Principles", 1)
    add_table(
        doc,
        ["Principle", "Meaning for this programme"],
        [
            ["Evidence before claims", "TRL 7 requires proof from an operational or pilot-like environment, not just code paths and local tests."],
            ["Stable foundations first", "Build features that remain valid even if AU ingestion changes: receipts, audit logs, access controls, state tracking, monitoring, and tests."],
            ["External integration by boundary", "If AU owns the webhook and calls the NEVERFLAT API, the evidence task is to prove that boundary, not to duplicate AU code."],
            ["Settlement must be verifiable", "A token spend should create a durable signed receipt that the frontend owner, EMP, or settlement service can verify and audit."],
        ],
        [1.9, 4.6],
    )

    add_heading(doc, "2. Phased Plan", 1)
    add_table(
        doc,
        ["Phase", "Focus", "Main outputs", "TRL 7 contribution"],
        [
            ["Phase 0", "Freeze scope and evidence assumptions", "Confirm deferred AU evidence; confirm receipt receiver; confirm pilot environment boundaries.", "Prevents documenting or building against a moving ingestion design."],
            ["Phase 1", "Signed spend receipts", "Receipt schema, signature, verification endpoint or helper, DB storage, API response, replay protection.", "Connects token spend to discount settlement proof."],
            ["Phase 2", "Auditability", "Append-only audit table for spend receipts, admin changes, wallet mode changes, awards, spends, failures.", "Shows traceability for reviewer and operational support."],
            ["Phase 3", "Pilot access controls", "No default admin credentials; scoped AU credential; manual test routes disabled in pilot; actor logging.", "Makes the pilot credible and reduces avoidable security criticism."],
            ["Phase 4", "Durable states and recovery", "Award/spend/receipt statuses, retry-required flags, confirmation states, failed-chain handling.", "Turns the system from demo flow into an operated service."],
            ["Phase 5", "Reconciliation and monitoring", "DB-vs-chain reconciliation job, health metrics, alerts, dashboards, transaction failure visibility.", "Supplies operational evidence for reliability."],
            ["Phase 6", "Evidence pack and hardening", "E2E tests, load/security tests, deployment runbook, smart contract artifacts, pilot screenshots/logs.", "Packages the claim for Horizon reporting."],
        ],
        [0.78, 1.25, 2.52, 1.95],
    )

    add_heading(doc, "3. Sprint 1: Start Here", 1)
    add_table(
        doc,
        ["Priority", "Task", "Acceptance check", "Owner dependency"],
        [
            ["1", "Define the spend receipt data model and signature format.", "Receipt includes ID, contract ID, wallet address, amount, session/provider IDs, chain ID, token contract, tx hash, status, timestamp, and backend signature.", "NEVERFLAT-owned."],
            ["2", "Create receipt persistence.", "Database table stores receipt payload, signature, status, linked spend tx, creation time, update time, and idempotency key.", "NEVERFLAT-owned."],
            ["3", "Return receipt from spend API.", "`/spend` and `/spend/me` return the signed receipt after successful token movement or reserved spend, depending on selected lifecycle.", "NEVERFLAT-owned."],
            ["4", "Add verification logic.", "A receiver can verify the signature and reject changed, replayed, expired, or mismatched receipts.", "Receiver integration can follow later."],
            ["5", "Add receipt tests.", "Tests cover valid receipt, tampered amount, wrong session ID, replay/idempotency, and invalid signature.", "NEVERFLAT-owned."],
        ],
        [0.7, 2.1, 2.55, 1.15],
    )

    add_callout(
        doc,
        "Sprint 1 decision",
        "Backend-to-backend receipt delivery is preferable for TRL 7, but frontend forwarding is acceptable if receipts are signed, durable, and verifiable. The first implementation should not depend on knowing the final receiver.",
    )

    add_heading(doc, "4. Workstream Detail", 1)
    add_heading(doc, "4.1 Signed Spend Receipt and Settlement Lifecycle", 2)
    add_para(
        doc,
        "Build a receipt as the business proof of spend. Token movement alone proves an on-chain transfer, but it does not prove that a charging discount was validly requested, linked to a session, accepted by the external system, and auditable later.",
    )
    add_table(
        doc,
        ["Item", "Implementation direction"],
        [
            ["Receipt statuses", "`reserved`, `settled`, `rejected`, `expired`, `cancelled`."],
            ["Signature", "Sign canonical JSON with a NEVERFLAT backend signing key. Publish verification rules and key identifier."],
            ["Delivery", "Return receipt to frontend immediately; optionally add server-to-server delivery once the receiver is agreed."],
            ["Replay prevention", "Use receipt ID and session/provider/idempotency key uniqueness."],
            ["Evidence", "Store signed receipt, verification result, token tx hash, and status transitions."],
        ],
        [1.55, 4.95],
    )

    add_heading(doc, "4.2 Audit Logs", 2)
    add_table(
        doc,
        ["Event family", "Minimum audit record"],
        [
            ["Receipt lifecycle", "Actor, receipt ID, previous status, new status, amount, session ID, provider ID, timestamp."],
            ["Admin changes", "Actor, rule version, old values, new values, timestamp, reason if supplied."],
            ["Wallet actions", "Actor, contract ID, wallet address, mode change or linked wallet action, signature result."],
            ["Award/spend failures", "Actor/system, stage, error, dedup key, tx hash if present, retry status."],
        ],
        [1.65, 4.85],
    )

    add_heading(doc, "4.3 Pilot Identity and RBAC", 2)
    add_table(
        doc,
        ["Control", "Pilot-ready target"],
        [
            ["Admin credentials", "No default admin password; credentials must be configured at deployment."],
            ["AU credential", "Dedicated credential scoped to ingestion only, once the AU integration shape is final."],
            ["User wallet access", "Pilot uses EMP-forwarded identity; manual contract ID lookup disabled outside test mode."],
            ["Admin permissions", "At minimum separate view-only from rule-editing permissions or document why a smaller pilot control is accepted."],
            ["Actor logging", "All privileged requests include actor identity in audit logs."],
        ],
        [1.55, 4.95],
    )

    add_heading(doc, "4.4 Durable Transaction States", 2)
    add_table(
        doc,
        ["Lifecycle object", "Recommended states"],
        [
            ["Award", "`accepted`, `eligible`, `submitted`, `confirmed`, `not_eligible`, `failed`, `retry_required`."],
            ["Spend", "`requested`, `validated`, `submitted`, `confirmed`, `failed`, `retry_required`."],
            ["Receipt", "`reserved`, `settled`, `rejected`, `expired`, `cancelled`."],
            ["Reconciliation", "`matched`, `missing_on_chain`, `missing_in_db`, `amount_mismatch`, `manual_review`."],
        ],
        [1.55, 4.95],
    )

    add_heading(doc, "5. Deferred Until Ingestion Stabilises", 1)
    add_table(
        doc,
        ["Deferred item", "Why defer", "What to capture later"],
        [
            ["AU evidence pack", "The ingestion flow may change, so documenting it now may create throwaway work.", "Final payload samples, request logs, accepted API responses, DB rows, tx hashes, and duplicate-call behaviour."],
            ["AU-specific monitoring", "Dashboards should follow the final operational boundary.", "Request rate, rejection rate, duplicate rate, latency, retry/failure rate, and alert thresholds."],
            ["AU payload mapping report", "Mapping should be frozen after AU payload shape is stable.", "Field-by-field mapping from AU CDR to NEVERFLAT canonical session."],
        ],
        [1.75, 2.2, 2.55],
    )

    add_heading(doc, "6. TRL 7 Acceptance Gates", 1)
    add_table(
        doc,
        ["Gate", "Pass condition"],
        [
            ["Functional gate", "Award, wallet update, spend, signed receipt, receipt verification, admin rules, and audit trail work end-to-end in a pilot environment."],
            ["Integration gate", "Final AU or provider ingestion path is evidenced with realistic CDRs and duplicate handling."],
            ["Operational gate", "Deployment, migrations, secrets, logs, monitoring, rollback, and support runbook are documented and exercised."],
            ["Blockchain gate", "Token contract address, chain, transaction hashes, confirmation handling, and reconciliation reports are available."],
            ["Security gate", "Pilot auth/RBAC controls are active; no default admin credential; sensitive operations are audited."],
            ["Testing gate", "Unit, API, DB integration, E2E, failure, and load tests have saved reports."],
        ],
        [1.55, 4.95],
    )

    add_heading(doc, "7. Horizon Evidence Checklist", 1)
    add_table(
        doc,
        ["Evidence item", "Target artifact"],
        [
            ["Pilot demo script", "Step-by-step script showing AU/pilot CDR, award, wallet balance, spend, signed receipt, and transaction history."],
            ["Screenshots", "Wallet UI, admin rule screen, transaction history, monitoring dashboard, and receipt verification result."],
            ["Logs", "Redacted API logs for CDR ingestion, spend request, receipt creation, and status changes."],
            ["Database extracts", "Awards, spends, receipts, audit logs, and reconciliation result rows with sensitive fields redacted."],
            ["Blockchain proof", "Polygon transaction hashes, token contract address, chain ID, and confirmation details."],
            ["Test reports", "Jest/API/E2E/load/security results with date, environment, version, and pass/fail summary."],
            ["Deployment record", "Image tag/digest, environment config checklist, migration logs, secrets method, and rollback notes."],
        ],
        [1.65, 4.85],
    )

    add_heading(doc, "8. Suggested Immediate Backlog", 1)
    add_table(
        doc,
        ["Order", "Backlog item", "Definition of done"],
        [
            ["1", "Create receipt schema and signing utility.", "Receipt can be generated and verified in automated tests."],
            ["2", "Add receipts migration and service methods.", "Receipt persistence supports uniqueness, status updates, and lookup by receipt ID/session ID."],
            ["3", "Wire receipt creation into spend endpoints.", "Spend response includes signed receipt; existing wallet balance behaviour remains intact."],
            ["4", "Add audit log migration and helper.", "Receipt, spend, admin, and wallet actions can emit audit events."],
            ["5", "Add tests and API docs.", "Build and tests pass; API docs show receipt fields and verification rules."],
            ["6", "Plan pilot RBAC changes.", "Manual routes, admin credentials, and AU scoped credential are specified for pilot deployment."],
        ],
        [0.65, 2.25, 3.6],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
